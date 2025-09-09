from xml.dom.minidom import Document
import pandas as pd
import numpy as np
import hvplot.pandas
import holoviews as hv
from holoviews import opts, dim


def cdf(df):
    """Return the cumulative distribution function of a dataframe."""
    df = df.apply(lambda x: x.sort_values().reset_index(drop=True), axis=0)
    df.index = np.linspace(0, 100, len(df))
    return df


# %%
def customize_legend(
    fig, legend_labels, legend_position="upper center", bbox_to_anchor=(0.5, 1.15)
):
    axes = fig.get_axes()[0]
    # put the legend outside the plot at the top centered
    axes.legend(
        legend_labels,
        loc=legend_position,
        bbox_to_anchor=bbox_to_anchor,
        ncols=len(legend_labels),
    )
    return fig


def rectangle_around_plot(fig, edgecolor="black", facecolor="none"):
    import matplotlib.patches as patches

    # fig.subplots_adjust(left=0.1, right=0.9, wspace=0.4, hspace=0.4)
    rect = patches.Rectangle(
        (0, 0.3),
        1,
        0.4,
        linewidth=1,
        edgecolor=edgecolor,
        facecolor=facecolor,
        transform=fig.transFigure,
        figure=fig,
    )

    # Add the patch to the figure
    fig.add_artist(rect)
    return fig


def exceedance_plot(
    df,
    ylabel,
    xlabel,
    line_styles=[
        (0, (5, 1, 2, 1)),
        "-",
        ":",
        "--",
        "-.",
        ":",
    ],
    legend_position="upper center",
):
    edf = cdf(df)
    line_plot = edf.hvplot.line(
        linestyle=line_styles,
        ylabel=ylabel,
        xlabel=xlabel,
        grid=True,
        legend="top",
    )
    line_plot.opts(fig_inches=5, show_frame=True).opts(opts.Layout(tight=True)).opts(
        backend_opts={
            "legend.frame_on": False,
        }
    )
    fig = hvplot.render(line_plot, backend="matplotlib")
    fig = customize_legend(fig, df.columns, legend_position=legend_position)
    fig = rectangle_around_plot(fig)
    return fig


def save_figure(fig, filename):
    fig.savefig(filename, dpi=300, bbox_inches="tight")


#
import pathlib


def save_to_word(
    word_file, image_file, image_alt_text=None, image_caption=None, image_width=6.0
):
    """
    Save an image to a Word document with metadata.

    word_file: The path to the Word document. If it exists, the image will be added to it, else a new one will be created.
    image_file: The path to the image file.
    image_alt_text: Alternative text for the image.
    image_caption: Caption for the image.
    image_width: Width of the image in inches.

    Note: requires pip install python-docx
    """
    try:
        from pydoc import doc
        from docx import Document
        from docx.shared import Inches
        from docx.oxml.shared import OxmlElement, qn
    except ImportError:
        print("Error: Required modules are not installed.")
        print("pip install python-docx")
        return
    if not pathlib.Path(word_file).exists():
        doc = Document()
    else:
        doc = Document(word_file)

    p = doc.add_paragraph()
    run = p.add_run()
    picture = run.add_picture(image_file, width=Inches(image_width))
    # Add alt text (good for accessibility/search inside Word)
    # (python-docx doesn’t expose alt text directly; set it via XML on the inline drawing)
    if image_alt_text:
        inline = picture._inline
        docPr = inline.docPr
        docPr.set("title", image_caption)
        docPr.set("descr", image_alt_text)
    if image_caption:
        caption = doc.add_paragraph(image_caption)
    doc.add_page_break()
    doc.save(word_file)
